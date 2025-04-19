# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog
import docx
import os
import copy
import re
from docx.shared import Cm
from docx.opc.exceptions import PackageNotFoundError
# Добавим импорты для проверки на изображения и стили
from docx.oxml.drawing import CT_Drawing
from docx.enum.style import WD_STYLE_TYPE
# --- НУЖНЫ ДЛЯ ВСТАВКИ ПРОБЕЛОВ ---
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# ------------------------------------

# --- Константы ---
HEADING_STYLES = ['Заголовок 1', 'Heading 1']
INVALID_FILENAME_CHARS = r'[\\/*?:"<>|]'
# --- НОВАЯ КОНСТАНТА: Количество пробелов для имитации отступа ---
# Установите нужное количество пробелов здесь (например, 2, 4, 5)
LEADING_SPACES_COUNT = 2
# -------------------------------------------------------------------
# --- НОВАЯ КОНСТАНТА: Текст-разделитель сцен ---
SCENE_SEPARATOR = "***"
# -------------------------------------------------------------------

# --- Функции (sanitize_filename, copy_run_formatting, copy_paragraph, copy_style_attributes, copy_document_styles - БЕЗ ИЗМЕНЕНИЙ) ---

def sanitize_filename(filename):
    """Очищает строку, чтобы она была допустимым именем файла."""
    sanitized = re.sub(INVALID_FILENAME_CHARS, '_', filename)
    sanitized = sanitized.strip('. ')
    if not sanitized:
        sanitized = "Без названия"
    return sanitized

def copy_run_formatting(source_run, target_run):
    """Копирует базовое форматирование из одного run в другой."""
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    if source_run.font.color and source_run.font.color.rgb:
         target_run.font.color.rgb = source_run.font.color.rgb

def copy_paragraph(source_para, target_doc):
    """
    Копирует параграф со всем форматированием в целевой документ.
    Возвращает вновь созданный параграф в target_doc.
    ПРИМЕЧАНИЕ: Эта версия НЕ копирует стили, если их нет в target_doc.
               Используйте более сложную версию при необходимости.
    """
    style_name = source_para.style.name if source_para.style else None
    target_style = None
    if style_name:
        try:
            target_style = target_doc.styles[style_name]
        except KeyError:
            pass # Будет использован стиль параграфа по умолчанию target_doc

    new_para = target_doc.add_paragraph(style=target_style)

    if new_para.runs:
        if not source_para.text.strip() and new_para.text.strip():
            for r in new_para.runs[:]:
                p = r._element.getparent()
                p.remove(r._element)
            new_para.text = ""

    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        copy_run_formatting(run, new_run)

    pf_target = new_para.paragraph_format
    pf_source = source_para.paragraph_format
    pf_target.alignment = pf_source.alignment
    pf_target.left_indent = pf_source.left_indent
    pf_target.right_indent = pf_source.right_indent
    pf_target.first_line_indent = pf_source.first_line_indent
    pf_target.space_before = pf_source.space_before
    pf_target.space_after = pf_source.space_after
    pf_target.line_spacing = pf_source.line_spacing
    pf_target.line_spacing_rule = pf_source.line_spacing_rule
    pf_target.keep_together = pf_source.keep_together
    pf_target.keep_with_next = pf_source.keep_with_next
    pf_target.page_break_before = pf_source.page_break_before
    pf_target.widow_control = pf_source.widow_control
    pf_target.tab_stops.clear_all()
    for ts in pf_source.tab_stops:
        pf_target.tab_stops.add_tab_stop(ts.position, ts.alignment, ts.leader)

    return new_para

def copy_style_attributes(source_style, target_style):
    """Копирует атрибуты из одного объекта стиля в другой (вспомогательная)."""
    try:
        target_style.base_style = source_style.base_style
        target_style.hidden = source_style.hidden
        target_style.locked = source_style.locked
        target_style.priority = source_style.priority
        target_style.quick_style = source_style.quick_style
        target_style.unhide_when_used = source_style.unhide_when_used

        if hasattr(source_style, 'font') and hasattr(target_style, 'font'):
             target_style.font.name = source_style.font.name
             target_style.font.size = source_style.font.size
             target_style.font.bold = source_style.font.bold
             target_style.font.italic = source_style.font.italic
             target_style.font.underline = source_style.font.underline
             if source_style.font.color and source_style.font.color.rgb:
                 target_style.font.color.rgb = source_style.font.color.rgb

        if hasattr(source_style, 'paragraph_format') and hasattr(target_style, 'paragraph_format'):
             pf_target = target_style.paragraph_format
             pf_source = source_style.paragraph_format
             pf_target.alignment = pf_source.alignment
             pf_target.first_line_indent = pf_source.first_line_indent
             pf_target.keep_together = pf_source.keep_together
             pf_target.keep_with_next = pf_source.keep_with_next
             pf_target.left_indent = pf_source.left_indent
             pf_target.line_spacing = pf_source.line_spacing
             pf_target.line_spacing_rule = pf_source.line_spacing_rule
             pf_target.page_break_before = pf_source.page_break_before
             pf_target.right_indent = pf_source.right_indent
             pf_target.space_after = pf_source.space_after
             pf_target.space_before = pf_source.space_before
             pf_target.widow_control = pf_source.widow_control
             pf_target.tab_stops.clear_all()
             for ts in pf_source.tab_stops:
                  pf_target.tab_stops.add_tab_stop(ts.position, ts.alignment, ts.leader)
    except Exception as e:
        print(f"Предупреждение при копировании атрибутов стиля '{source_style.name}': {e}")

def copy_document_styles(source_doc, target_doc):
    """Копирует все стили из одного документа в другой."""
    print("Копирование стилей...")
    num_styles_copied = 0
    num_styles_existing = 0
    num_styles_failed = 0
    existing_style_names = [s.name for s in target_doc.styles]

    for style in source_doc.styles:
        if style.name in existing_style_names:
            num_styles_existing += 1
            continue
        try:
            style_type = style.type
            if not isinstance(style_type, WD_STYLE_TYPE):
                if isinstance(style, docx.styles.style._ParagraphStyle):
                    style_type = WD_STYLE_TYPE.PARAGRAPH
                elif isinstance(style, docx.styles.style._CharacterStyle):
                    style_type = WD_STYLE_TYPE.CHARACTER
                elif isinstance(style, docx.styles.style._TableStyle):
                    style_type = WD_STYLE_TYPE.TABLE
                elif isinstance(style, docx.styles.style._NumberingStyle):
                     style_type = WD_STYLE_TYPE.LIST
                else:
                     print(f"Предупреждение: Неизвестный тип стиля '{style.name}' ({type(style)}). Пропуск.")
                     num_styles_failed += 1
                     continue
            target_style = target_doc.styles.add_style(style.name, style_type, style.builtin)
            copy_style_attributes(style, target_style)
            num_styles_copied += 1
        except ValueError as e:
             if "style name" in str(e).lower() and "already exists" in str(e).lower():
                  num_styles_existing += 1
             else:
                print(f"Предупреждение ValueError при копировании стиля '{style.name}': {e}")
                num_styles_failed += 1
        except Exception as e:
             print(f"Ошибка при копировании стиля '{style.name}': {e}")
             num_styles_failed += 1
    print(f"Копирование стилей завершено: Скопировано={num_styles_copied}, Существовало={num_styles_existing}, Ошибок={num_styles_failed}")


# --- ИЗМЕНЕННАЯ ФУНКЦИЯ ---
def apply_formatting_rules(paragraph):
    """
    Применяет правила форматирования к параграфу:
    - Диалоги ('—'): Убирает отступ первой строки.
    - Заголовки: Оставляет исходный отступ (из стиля/копирования).
    - Обычный текст: Убирает отступ первой строки и добавляет LEADING_SPACES_COUNT пробелов в начало.
    - НЕ ТРОГАЕТ параграфы, являющиеся разделителем сцен (SCENE_SEPARATOR).
    """
    paragraph_text_stripped = paragraph.text.strip()
    is_heading = paragraph.style and paragraph.style.name in HEADING_STYLES

    # Если это разделитель сцен, ничего не делаем с форматированием
    if paragraph_text_stripped == SCENE_SEPARATOR:
        return

    if paragraph_text_stripped.startswith("—"):
        # Диалог: Убираем любой отступ первой строки
        paragraph.paragraph_format.first_line_indent = Cm(0)
    elif is_heading:
        # Заголовок: Ничего не делаем, оставляем как есть (скопированный/стилевой отступ)
        pass
    else:
        # Обычный текст: Убираем отступ и добавляем пробелы
        paragraph.paragraph_format.first_line_indent = Cm(0) # Убираем визульный отступ

        if LEADING_SPACES_COUNT > 0:
            # Добавляем пробелы в начало параграфа через XML
            spaces = " " * LEADING_SPACES_COUNT
            try:
                # Создаем XML элементы для run (w:r) и text (w:t)
                run_element = OxmlElement('w:r')
                text_element = OxmlElement('w:t')
                # Важно: xml:space="preserve" говорит Word сохранять пробелы
                text_element.set(qn('xml:space'), 'preserve')
                text_element.text = spaces
                run_element.append(text_element)

                # Получаем XML-элемент параграфа (w:p)
                p_element = paragraph._element
                # Находим элемент свойств параграфа (w:pPr), если он есть
                pPr = p_element.xpath('./w:pPr')
                if pPr:
                    # Вставляем наш run с пробелами *после* свойств параграфа,
                    # но перед остальным содержимым
                    pPr[0].addnext(run_element)
                else:
                    # Если свойств параграфа нет, вставляем run в самое начало
                    p_element.insert(0, run_element)
            except Exception as xml_err:
                print(f"Предупреждение: Не удалось добавить пробелы в параграф: {xml_err}")
                # В качестве запасного варианта, можно попробовать добавить пробелы к первому run,
                # но это может нарушить форматирование, если первый run был, например, жирным.
                # if paragraph.runs:
                #     paragraph.runs[0].text = spaces + paragraph.runs[0].text


# --- ИЗМЕНЕННАЯ ФУНКЦИЯ ---
def remove_empty_paragraphs_and_format(input_path, output_path, progress_callback=None):
    """
    Открывает Word документ, удаляет пустые параграфы, копирует стили,
    применяет правила форматирования (пробелы вместо отступа),
    добавляет пустые строки вокруг SCENE_SEPARATOR и сохраняет результат.
    """
    try:
        print("Открытие исходного документа...")
        source_doc = docx.Document(input_path)
        target_doc = docx.Document()
        if target_doc.paragraphs:
             if not target_doc.paragraphs[0].text.strip() and not target_doc.paragraphs[0].runs:
                 p = target_doc.paragraphs[0]._element
                 p.getparent().remove(p)

        copy_document_styles(source_doc, target_doc)

        print("Обработка контента основного документа...")
        total = len(source_doc.element.body)
        for i, element in enumerate(source_doc.element.body):
            if progress_callback and total > 0:
                progress_callback(i / total)
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(element, source_doc)
                paragraph_text_stripped = para.text.strip()
                is_empty_para = not paragraph_text_stripped and not any(isinstance(run._element, CT_Drawing) for run in para.runs)
                if is_empty_para:
                    continue

                # --- НАЧАЛО ИЗМЕНЕНИЯ ДЛЯ SCENE_SEPARATOR ---
                if paragraph_text_stripped == SCENE_SEPARATOR:
                    # Добавляем пустую строку ПЕРЕД разделителем
                    target_doc.add_paragraph("")
                    # Копируем сам параграф-разделитель
                    copy_paragraph(para, target_doc) # Просто копируем, без apply_formatting_rules
                    # Добавляем пустую строку ПОСЛЕ разделителя
                    target_doc.add_paragraph("")
                else:
                    # Обрабатываем все остальные непустые параграфы как раньше
                    new_para = copy_paragraph(para, target_doc)
                    # Применяем НОВЫЕ правила форматирования (с пробелами),
                    # которые теперь игнорируют SCENE_SEPARATOR
                    apply_formatting_rules(new_para)
                # --- КОНЕЦ ИЗМЕНЕНИЯ ДЛЯ SCENE_SEPARATOR ---

            elif isinstance(element, docx.oxml.table.CT_Tbl):
                new_element = copy.deepcopy(element)
                target_doc.element.body.append(new_element)
            else:
                 try:
                      new_element = copy.deepcopy(element)
                      target_doc.element.body.append(new_element)
                 except Exception as copy_err:
                      print(f"Предупреждение: Не удалось скопировать элемент типа {type(element)}. Ошибка: {copy_err}")

        print("Сохранение основного документа...")
        target_doc.save(output_path)
        print(f"\n--- Обработка основного файла ЗАВЕРШЕНА ---")
        print(f"Успешно! Пустые параграфы удалены.")
        print(f"Правила форматирования (пробелы вместо отступа, отступы для '{SCENE_SEPARATOR}') применены.")
        print(f"Обработанный документ сохранен как: {output_path}")
        return True

    except PackageNotFoundError:
        print(f"Ошибка: Файл '{input_path}' не найден или не является корректным файлом DOCX.")
        return False
    except FileNotFoundError:
        print(f"Ошибка: Файл не найден по пути {input_path}")
        return False
    except ImportError:
        print("Ошибка: Библиотека python-docx не установлена.")
        return False
    except Exception as e:
        print(f"\nПроизошла ошибка при обработке основного документа:")
        print(e)
        import traceback
        traceback.print_exc()
        return False


# --- ИЗМЕНЕННАЯ ФУНКЦИЯ ---
def extract_chapters(input_path, chapters_output_dir, progress_callback=None):
    """
    Извлекает главы, копирует стили, применяет правила форматирования (пробелы вместо отступа)
    и добавляет пустые строки вокруг SCENE_SEPARATOR.
    """
    print(f"\n--- Извлечение глав ---")
    try:
        print("Открытие исходного документа для извлечения глав...")
        source_doc = docx.Document(input_path)
        os.makedirs(chapters_output_dir, exist_ok=True)
        print(f"Папка для глав: {chapters_output_dir}")

        # --- Определяем количество глав (заголовков) для прогресса ---
        headings = []
        for element in source_doc.element.body:
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(element, source_doc)
                if para.style and para.style.name in HEADING_STYLES and para.text.strip():
                    headings.append(para.text.strip())
        total_chapters = len(headings) if headings else 1
        current_chapter_idx = 0

        current_chapter_doc = None
        current_chapter_title = None
        chapter_count = 0

        for element in source_doc.element.body:
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(element, source_doc)
                is_heading = para.style and para.style.name in HEADING_STYLES
                paragraph_text_stripped = para.text.strip()
                is_empty_para = not paragraph_text_stripped and not any(isinstance(run._element, CT_Drawing) for run in para.runs)

                if is_heading and not paragraph_text_stripped:
                    print(f"Предупреждение: Найден пустой параграф со стилем заголовка. Пропуск.")
                    continue

                if is_heading:
                    # Завершаем предыдущую главу
                    if current_chapter_doc and current_chapter_title:
                        chapter_filename = sanitize_filename(current_chapter_title) + ".docx"
                        chapter_filepath = os.path.join(chapters_output_dir, chapter_filename)
                        try:
                            print(f"Сохранение главы: {chapter_filename}...")
                            current_chapter_doc.save(chapter_filepath)
                            print(f"Сохранена глава: {chapter_filename}")
                            chapter_count += 1
                            # --- Обновляем прогресс ---
                            if progress_callback and total_chapters > 0:
                                current_chapter_idx += 1
                                progress_callback(current_chapter_idx / total_chapters)
                        except Exception as save_err:
                             print(f"Ошибка сохранения главы '{chapter_filename}': {save_err}")

                    # Начинаем новую главу
                    current_chapter_title = paragraph_text_stripped
                    print(f"\nНачало обработки главы: '{current_chapter_title}'")
                    current_chapter_doc = docx.Document()
                    if current_chapter_doc.paragraphs:
                        if not current_chapter_doc.paragraphs[0].text.strip() and not current_chapter_doc.paragraphs[0].runs:
                            p = current_chapter_doc.paragraphs[0]._element
                            p.getparent().remove(p)

                    copy_document_styles(source_doc, current_chapter_doc)
                    # Копируем сам заголовок в новую главу
                    new_heading_para = copy_paragraph(para, current_chapter_doc)
                    # Применяем правила к заголовку (функция сама поймет, что делать)
                    apply_formatting_rules(new_heading_para) # Заголовки не меняются функцией

                elif current_chapter_doc and not is_empty_para: # Копируем только непустые параграфы внутри главы
                    # --- НАЧАЛО ИЗМЕНЕНИЯ ДЛЯ SCENE_SEPARATOR В ГЛАВАХ ---
                    if paragraph_text_stripped == SCENE_SEPARATOR:
                        # Добавляем пустую строку ПЕРЕД разделителем
                        current_chapter_doc.add_paragraph("")
                        # Копируем сам параграф-разделитель
                        copy_paragraph(para, current_chapter_doc) # Просто копируем
                        # Добавляем пустую строку ПОСЛЕ разделителя
                        current_chapter_doc.add_paragraph("")
                    else:
                        # Обрабатываем все остальные непустые параграфы как раньше
                        new_para_in_chapter = copy_paragraph(para, current_chapter_doc)
                        # Применяем НОВЫЕ правила форматирования (с пробелами),
                        # которые теперь игнорируют SCENE_SEPARATOR
                        apply_formatting_rules(new_para_in_chapter)
                    # --- КОНЕЦ ИЗМЕНЕНИЯ ДЛЯ SCENE_SEPARATOR В ГЛАВАХ ---

            elif isinstance(element, docx.oxml.table.CT_Tbl) and current_chapter_doc:
                 new_element = copy.deepcopy(element)
                 current_chapter_doc.element.body.append(new_element)

            elif current_chapter_doc: # Копируем прочие элементы (например, картинки в параграфах, которые не прошли проверку выше, но нужно скопировать)
                 try:
                      # Проверяем, не пустой ли это параграф (еще раз, на всякий случай, если логика пропуска выше не сработала для сложных случаев)
                      if isinstance(element, docx.oxml.text.paragraph.CT_P):
                          para_check = docx.text.paragraph.Paragraph(element, source_doc)
                          is_really_empty = not para_check.text.strip() and not any(isinstance(run._element, CT_Drawing) for run in para_check.runs)
                          if is_really_empty:
                              continue # Пропускаем явно пустые

                      new_element = copy.deepcopy(element)
                      current_chapter_doc.element.body.append(new_element)
                 except Exception as copy_err:
                      print(f"Предупреждение (глава '{current_chapter_title}'): Не удалось скопировать элемент типа {type(element)}. Ошибка: {copy_err}")


        # Сохраняем последнюю главу
        if current_chapter_doc and current_chapter_title:
            chapter_filename = sanitize_filename(current_chapter_title) + ".docx"
            chapter_filepath = os.path.join(chapters_output_dir, chapter_filename)
            try:
                print(f"Сохранение последней главы: {chapter_filename}...")
                current_chapter_doc.save(chapter_filepath)
                print(f"Сохранена глава: {chapter_filename}")
                chapter_count += 1
            except Exception as save_err:
                 print(f"Ошибка сохранения последней главы '{chapter_filename}': {save_err}")

        if chapter_count > 0:
            print(f"\nУспешно извлечено глав: {chapter_count}")
        else:
            print("\nНе найдено параграфов со стилями заголовков глав или главы не были сохранены.")
            print(f"Проверьте, что в документе используются стили: {', '.join(HEADING_STYLES)}")
            try:
                # Удаляем пустую папку глав, если она была создана, но ни одна глава не сохранена
                if os.path.exists(chapters_output_dir) and not os.listdir(chapters_output_dir):
                    os.rmdir(chapters_output_dir)
                    print(f"Пустая папка для глав удалена: {chapters_output_dir}")
            except OSError as e:
                print(f"Не удалось удалить папку для глав {chapters_output_dir}: {e}")

    except PackageNotFoundError:
        print(f"Ошибка: Файл '{input_path}' не найден или не является корректным файлом DOCX.")
    except FileNotFoundError:
        print(f"Ошибка: Файл не найден по пути {input_path}")
    except ImportError:
        print("Ошибка: Библиотека python-docx не установлена.")
    except Exception as e:
        print(f"\nПроизошла ошибка при извлечении глав:")
        print(e)
        import traceback
        traceback.print_exc()

# --- Основной блок и GUI ---
def process_document_gui():
    import customtkinter as ctk
    input_file_path = filedialog.askopenfilename(
        title="Выберите исходный Word документ",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not input_file_path:
        return
    base_name = os.path.basename(input_file_path)
    name, ext = os.path.splitext(base_name)
    default_save_name = f"{name}_обработанный{ext}"
    initial_dir = os.path.dirname(input_file_path)
    output_file_path = filedialog.asksaveasfilename(
        title="Сохранить обработанный файл как...",
        initialdir=initial_dir,
        initialfile=default_save_name,
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not output_file_path:
        return
    if not output_file_path.lower().endswith('.docx'):
        output_file_path += '.docx'
    # --- Окно прогресса ---
    progress_win = ctk.CTkToplevel()
    progress_win.title("Обработка документа")
    # Центрируем окно
    win_width, win_height = 420, 110
    screen_width = progress_win.winfo_screenwidth()
    screen_height = progress_win.winfo_screenheight()
    x = (screen_width // 2) - (win_width // 2)
    y = (screen_height // 2) - (win_height // 2)
    progress_win.geometry(f"{win_width}x{win_height}+{x}+{y}")
    # Делаем окно поверх основного
    if progress_win.master:
        progress_win.transient(progress_win.master)
    progress_win.lift()
    progress_win.attributes('-topmost', True)
    label = ctk.CTkLabel(progress_win, text="Обработка документа...", font=("Arial", 13))
    label.pack(pady=(18, 8))
    progress_bar = ctk.CTkProgressBar(progress_win, width=340, height=18)
    progress_bar.pack(pady=(0, 8))
    progress_bar.set(0)
    progress_win.update()
    def on_progress(value):
        progress_bar.set(value)
        progress_win.update()
    success = remove_empty_paragraphs_and_format(input_file_path, output_file_path, on_progress)
    progress_win.destroy()
    if success:
        def open_folder():
            os.startfile(os.path.dirname(output_file_path))
        win = ctk.CTkToplevel()
        win.title("Успех")
        # Центрируем окно
        win_width, win_height = 460, 170
        win.update_idletasks()
        screen_width = win.winfo_screenwidth()
        screen_height = win.winfo_screenheight()
        x = (screen_width // 2) - (win_width // 2)
        y = (screen_height // 2) - (win_height // 2)
        win.geometry(f"{win_width}x{win_height}+{x}+{y}")
        ctk.CTkLabel(win, text=f"Документ обработан и сохранён как:\n{output_file_path}", wraplength=320).pack(pady=10)
        ctk.CTkButton(win, text="Открыть папку", command=open_folder).pack(pady=5)
        ctk.CTkButton(win, text="Закрыть", command=win.destroy).pack(pady=5)
        win.after(500, lambda: win.focus_force())
    else:
        messagebox.showerror("Ошибка", "Произошла ошибка при обработке документа.")

def extract_chapters_gui():
    import customtkinter as ctk
    input_file_path = filedialog.askopenfilename(
        title="Выберите исходный Word документ для извлечения глав",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not input_file_path:
        return
    # Формируем папку по умолчанию
    base_name = os.path.basename(input_file_path)
    name, _ = os.path.splitext(base_name)
    default_dir = os.path.join(os.path.dirname(input_file_path), f"{name}_Главы")
    if not os.path.exists(default_dir):
        os.makedirs(default_dir)
    use_default = messagebox.askyesno("Папка для глав", f"Использовать папку по умолчанию для сохранения глав?\n{default_dir}")
    if use_default:
        output_dir = default_dir
    else:
        output_dir = filedialog.askdirectory(title="Выберите папку для сохранения глав")
        if not output_dir:
            return
    # --- Окно прогресса ---
    progress_win = ctk.CTkToplevel()
    progress_win.title("Извлечение глав")
    label = ctk.CTkLabel(progress_win, text="Извлечение и обработка глав...", font=("Arial", 13))
    label.pack(pady=(18, 8))
    progress_bar = ctk.CTkProgressBar(progress_win, width=340, height=18)
    progress_bar.pack(pady=(0, 8))
    progress_bar.set(0)
    progress_win.update_idletasks()
    # Центрируем окно
    win_width, win_height = 420, 110
    screen_width = progress_win.winfo_screenwidth()
    screen_height = progress_win.winfo_screenheight()
    x = (screen_width // 2) - (win_width // 2)
    y = (screen_height // 2) - (win_height // 2)
    progress_win.geometry(f"{win_width}x{win_height}+{x}+{y}")
    # Делаем окно поверх основного
    if progress_win.master:
        progress_win.transient(progress_win.master)
    progress_win.lift()
    progress_win.attributes('-topmost', True)
    progress_win.update()
    def on_progress(value):
        progress_bar.set(value)
        progress_win.update()
    try:
        extract_chapters(input_file_path, output_dir, on_progress)
        progress_win.destroy()
        def open_folder():
            os.startfile(output_dir)
        win = ctk.CTkToplevel()
        win.title("Успех")
        # Центрируем окно
        win_width, win_height = 460, 170
        win.update_idletasks()
        screen_width = win.winfo_screenwidth()
        screen_height = win.winfo_screenheight()
        x = (screen_width // 2) - (win_width // 2)
        y = (screen_height // 2) - (win_height // 2)
        win.geometry(f"{win_width}x{win_height}+{x}+{y}")
        ctk.CTkLabel(win, text=f"Главы успешно извлечены в папку:\n{output_dir}", wraplength=320).pack(pady=10)
        ctk.CTkButton(win, text="Открыть папку", command=open_folder).pack(pady=5)
        ctk.CTkButton(win, text="Закрыть", command=win.destroy).pack(pady=5)
        win.after(500, lambda: win.focus_force())
    except Exception as e:
        progress_win.destroy()
        messagebox.showerror("Ошибка", f"Произошла ошибка при извлечении глав:\n{e}")

if __name__ == "__main__":
    import customtkinter as ctk
    import tkinter.messagebox as messagebox
    ctk.set_appearance_mode("System")  # Можно выбрать "Dark" или "Light"
    ctk.set_default_color_theme("blue")

    class FicbookApp(ctk.CTk):
        def __init__(self):
            super().__init__()
            self.title("Ficbook DOCX Tools")
            self.geometry("400x240")
            self.resizable(False, False)
            # Центрирование окна на экране
            self.update_idletasks()
            width = 400
            height = 240
            x = (self.winfo_screenwidth() // 2) - (width // 2)
            y = (self.winfo_screenheight() // 2) - (height // 2)
            self.geometry(f"{width}x{height}+{x}+{y}")
            # Центральный фрейм
            frame = ctk.CTkFrame(self, corner_radius=15)
            frame.pack(padx=20, pady=20, fill="both", expand=True)
            label = ctk.CTkLabel(frame, text="Выберите действие:", font=("Arial", 16, "bold"))
            label.pack(pady=(10, 20))
            btn1 = ctk.CTkButton(frame, text="Обработать документ", width=220, height=40, font=("Arial", 13), command=process_document_gui)
            btn1.pack(pady=10)
            btn2 = ctk.CTkButton(frame, text="Извлечь и обработать главы", width=220, height=40, font=("Arial", 13), command=extract_chapters_gui)
            btn2.pack(pady=10)
            # Можно добавить иконку: self.iconbitmap('your_icon.ico')
    app = FicbookApp()
    app.mainloop()